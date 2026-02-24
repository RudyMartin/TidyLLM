import streamlit as st
from datetime import date
import io
from typing import List, Tuple

import psycopg2
import boto3
from pypdf import PdfReader
from docx import Document

from app.config import settings, get_run_prefix
from workers.standards_loader import load_standards_from_s3
from workers.summarizer import compress_chunks, build_section_md
from workers.checker import run_checks_s3
from app.db_utils import save_run
from app.metrics import Metrics

# ---- AWS S3 client ----
S3 = boto3.client("s3")

st.set_page_config(page_title="QA Validation Review – Runner", layout="wide")
st.title("QA Validation Review – Runner")

# =========================
# 0) Identity & Scope
# =========================
user = st.selectbox("Choose your name", ["alex", "maria", "ravi"], index=0)

ORG_MAP = {
    "QA-Americas": {"teams": ["QA-East", "QA-West"], "processes": ["QAValidationReview"]},
    "QA-EMEA": {"teams": ["QA-UK", "QA-DE", "QA-FR"], "processes": ["QAValidationReview"]},
    "QA-APAC": {"teams": ["QA-SG", "QA-AU"], "processes": ["QAValidationReview"]},
}
org = st.sidebar.selectbox("Org", list(ORG_MAP.keys()), index=0)
team = st.sidebar.selectbox("Team", ORG_MAP[org]["teams"], index=0)
process = st.sidebar.selectbox("Process", ORG_MAP[org]["processes"], index=0)
review_date = st.sidebar.date_input("Review date (point-in-time standards)", value=date.today())

# Per-user+team run folder in S3 (no local files)
DOC_ID = "WF-2025-042"  # demo id (stable for comparisons)
doc_prefix = f"{get_run_prefix(user, team)}/{DOC_ID}"
st.caption(f"Run folder: s3://{settings.bucket}/{doc_prefix}")

# Burger counter (team/process scoped)
def _scalar(q: str, params: tuple):
    try:
        conn = psycopg2.connect(settings.pg_dsn)
        cur = conn.cursor(); cur.execute(q, params)
        x = cur.fetchone()[0]
        conn.close()
        return x or 0
    except Exception:
        return 0

burgers = _scalar(
    """
    SELECT COALESCE(SUM(finding_count),0)
    FROM review_runs
    WHERE (%s IS NULL OR org=%s) AND (%s IS NULL OR team=%s) AND (%s IS NULL OR process=%s)
    """,
    (org, org, team, team, process, process)
)
st.metric("🍔 Findings Checked (Burgers Served)", f"{burgers:,}")

st.divider()

# =========================
# 1) Standards (S3 → atoms)
# =========================
st.sidebar.subheader("Standards")
standards_prefix = st.sidebar.text_input("S3 standards prefix", value="standards/qa/")
standards_eff = st.sidebar.date_input("Standards effective date", value=date.today())
if st.sidebar.button("Load Standards from S3 → Atoms"):
    atoms = load_standards_from_s3(settings.bucket, standards_prefix, stage=process, eff_date=standards_eff)
    count, preview_key = 0, None
    for rel_key, md in atoms:
        s3_key = f"{doc_prefix}/requirements/{rel_key}"
        S3.put_object(Bucket=settings.bucket, Key=s3_key, Body=md.encode("utf-8"), ContentType="text/markdown")
        count += 1
        if preview_key is None: preview_key = s3_key
    st.success(f"Created {count} rule atoms under {doc_prefix}/requirements/atoms/")
    if preview_key:
        body = S3.get_object(Bucket=settings.bucket, Key=preview_key)["Body"].read().decode("utf-8", errors="ignore")
        st.code(body[:1200], language="markdown")

# =========================
# 2) Uploads → in-memory text
# =========================
st.subheader("1) Upload documents")
col1, col2 = st.columns(2)
review_file = col1.file_uploader("Validation Review (PDF/DOCX)", type=["pdf","docx"])
scope_file  = col2.file_uploader("Validation Scope (PDF/DOCX)", type=["pdf","docx"])

def read_pdf_bytes(b: bytes) -> str:
    out = []
    reader = PdfReader(io.BytesIO(b))
    for p in reader.pages:
        t = p.extract_text() or ""
        if t: out.append(t)
    return "\n\n".join(out)

def read_docx_bytes(b: bytes) -> str:
    doc = Document(io.BytesIO(b))
    parts = [p.text for p in doc.paragraphs if p.text]
    for tbl in doc.tables:
        for row in tbl.rows:
            row_txt = " | ".join(c.text for c in row.cells).strip()
            if row_txt: parts.append(row_txt)
    return "\n\n".join(parts)

def extract_text(upload) -> Tuple[str, str]:
    if upload is None:
        return "", ""
    name = upload.name
    ext = name.lower().split(".")[-1]
    b = upload.read()
    if ext == "pdf":
        return name, read_pdf_bytes(b)
    if ext == "docx":
        return name, read_docx_bytes(b)
    return name, ""

st.subheader("2) Ingest → Chunks")
chunk_count_placeholder = st.empty()

if st.button("Create Chunks from Uploads"):
    if not review_file and not scope_file:
        st.error("Please upload at least the Validation Review file.")
    else:
        texts: List[str] = []
        evidence_links: List[str] = []
        for up in [review_file, scope_file]:
            if not up: continue
            fname, txt = extract_text(up)
            if not txt: continue
            texts.append(f"# Source: {fname}\n\n{txt}")
            raw_key = f"{doc_prefix}/00_source/{fname}"
            S3.put_object(Bucket=settings.bucket, Key=raw_key, Body=up.getvalue())
            evidence_links.append(f"s3://{settings.bucket}/{raw_key}")
        
        full_text = "\n\n".join(texts)

        # Simple chunker: split by ~1200 chars on paragraph boundaries
        def chunks(s: str, maxlen=1200):
            parts, cur = [], []
            for para in s.split("\n\n"):
                if sum(len(p) for p in cur)+len(para) > maxlen and cur:
                    parts.append("\n\n".join(cur)); cur=[]
                cur.append(para)
            if cur: parts.append("\n\n".join(cur))
            return parts

        md_chunks: List[str] = []
        for i, c in enumerate(chunks(full_text), start=1):
            md = f"""---
source: upload
index: {i}
---
{c}
"""
            key = f"{doc_prefix}/10_chunks/chunk_{i:04d}.md"
            S3.put_object(Bucket=settings.bucket, Key=key, Body=md.encode("utf-8"), ContentType="text/markdown")
            md_chunks.append(key)
        st.session_state["chunks"] = md_chunks
        st.session_state["evidence"] = evidence_links
        chunk_count_placeholder.success(f"Chunks created: {len(md_chunks)} ✅")

# =========================
# 3) Summarize → Section
# =========================
st.subheader("3) Summarize → Section")

if st.button("Summarize Chunks"):
    keys = st.session_state.get("chunks", [])
    if not keys:
        st.error("Create chunks first.")
    else:
        md_blobs = [S3.get_object(Bucket=settings.bucket, Key=k)["Body"].read().decode("utf-8", errors="ignore") for k in keys]
        body = compress_chunks(md_blobs, max_chars=8*settings.max_gist_tokens)
        chunk_ids = [k.split("/")[-1].replace(".md", "") for k in keys]
        section_md = build_section_md(DOC_ID, "3.2", chunk_ids, st.session_state.get("evidence", []), body)
        out_key = f"{doc_prefix}/20_sections/3.2.section.md"
        S3.put_object(Bucket=settings.bucket, Key=out_key, Body=section_md.encode("utf-8"), ContentType="text/markdown")
        st.session_state["section_key"] = out_key
        st.success("Section summary ready ✅")
        st.code(section_md[:1200], language="markdown")

# =========================
# 4) Check → Findings (S3-native rules)
# =========================
st.subheader("4) Check → Findings")
if st.button("Run Checks"):
    sk = st.session_state.get("section_key")
    if not sk:
        st.error("Summarize first.")
    else:
        section_md = S3.get_object(Bucket=settings.bucket, Key=sk)["Body"].read().decode("utf-8", errors="ignore")
        header, body_md = section_md.split("---\n# Section Gist", 1)
        evidence = [line.strip()[1:].strip() for line in header.splitlines() if line.strip().startswith("-")]
        rules_prefix = f"{doc_prefix}/requirements/atoms"
        finding_mds, ruleset_hash = run_checks_s3(body_md, evidence, settings.bucket, rules_prefix, as_of=review_date)
        keys = []
        for i, fmd in enumerate(finding_mds):
            out_key = f"{doc_prefix}/30_findings/3.2.finding.{i:02d}.mf.md"
            S3.put_object(Bucket=settings.bucket, Key=out_key, Body=fmd.encode("utf-8"), ContentType="text/markdown")
            keys.append(out_key)
        st.session_state["finding_keys"] = keys
        st.session_state["ruleset_hash"] = ruleset_hash
        st.success(f"Findings generated: {len(keys)} (hash {ruleset_hash[:8]}…) ✅")

# =========================
# 5) Assemble → Report
# =========================
st.subheader("5) Assemble → Report")
if st.button("Assemble Report"):
    sk = st.session_state.get("section_key")
    fks = st.session_state.get("finding_keys", [])
    if not sk or not fks:
        st.error("Run summary and checks first.")
    else:
        section_md = S3.get_object(Bucket=settings.bucket, Key=sk)["Body"].read().decode("utf-8", errors="ignore")
        findings_md = [S3.get_object(Bucket=settings.bucket, Key=k)["Body"].read().decode("utf-8", errors="ignore") for k in fks]
        report = "\n\n".join([
            """---
title: QA Validation Review Report
---""",
            "# Executive Summary\n- QA Validation Review demo run.",
            "# Section\n" + section_md,
            "# Findings\n" + "\n\n".join(findings_md)
        ])
        out_key = f"{doc_prefix}/90_report/final.md"
        S3.put_object(Bucket=settings.bucket, Key=out_key, Body=report.encode("utf-8"), ContentType="text/markdown")
        st.session_state["report_key"] = out_key
        st.success("Final report saved to S3 ✅")
        st.code(report[:1200], language="markdown")

# =========================
# 6) Save Run to DB
# =========================
st.subheader("6) Save Run to DB")
if st.button("💾 Save Run to DB"):
    try:
        finding_keys = st.session_state.get("finding_keys", [])
        finding_mds  = [S3.get_object(Bucket=settings.bucket, Key=k)["Body"].read().decode("utf-8", errors="ignore") for k in finding_keys]
        # minimal metrics snapshot
        m = Metrics(pipeline_version="v1.2")
        m.add_evidence(len(st.session_state.get("evidence", [])))
        m_snapshot = m.snapshot()
        conn = psycopg2.connect(settings.pg_dsn)
        run_id = save_run(
            conn,
            reviewer=user,
            doc_id=DOC_ID,
            review_date=review_date,
            ruleset_hash=st.session_state.get("ruleset_hash"),
            report_s3_uri=f"s3://{settings.bucket}/{st.session_state.get('report_key')}",
            finding_mds=finding_mds,
            metrics=m_snapshot,
            org=org, team=team, process=process,
        )
        conn.close()
        st.success(f"Saved review run #{run_id} ✅")
        st.balloons()
    except Exception as e:
        st.error(f"DB save failed: {e}")

# =========================
# Help
# =========================
with st.expander("Help"):
    st.markdown("""
**Flow**
1) Load Standards → Atoms (optional if already generated)
2) Upload Review (+ Scope)
3) Create Chunks → Summarize → Run Checks → Assemble Report
4) Save Run to DB → burger counter ticks up

**Team-based S3 Storage**: Your files are stored in `s3://bucket/{team}/{user}/WF-2025-042/`
""")