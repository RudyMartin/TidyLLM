"""
Text Extraction Implementation

Extract text content from various document formats with support for:
- PDF documents (configurable page limits)
- Microsoft Word documents (DOCX)
- Plain text files
- Error handling and fallback options

Part of the tidyllm-verse: Educational ML with complete transparency
"""

import os
from pathlib import Path
from typing import Optional, Tuple

# Try to import document processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True  
except ImportError:
    DOCX_AVAILABLE = False

class TextExtractor:
    """Extract text content from various document formats."""
    
    def __init__(self):
        """Initialize text extractor with available libraries."""
        self.pdf_available = PDF_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
        
        if not PDF_AVAILABLE:
            print("Info: PyPDF2 not available. PDF extraction disabled. Install with: pip install PyPDF2")
        if not DOCX_AVAILABLE:
            print("Info: python-docx not available. DOCX extraction disabled. Install with: pip install python-docx")
    
    def extract_text(self, file_path: str, max_pages: int = 5) -> Tuple[str, dict]:
        """
        Extract text from document with metadata about extraction.
        
        Args:
            file_path: Path to document file
            max_pages: Maximum pages to extract (for PDF/DOCX)
            
        Returns:
            Tuple of (extracted_text, extraction_metadata)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return "", {"error": "File not found", "file_path": str(file_path)}
        
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_path, max_pages)
            elif file_extension == '.docx':
                return self._extract_docx_text(file_path, max_pages)  
            elif file_extension in ['.txt', '.text']:
                return self._extract_text_file(file_path)
            else:
                return "", {"error": f"Unsupported file type: {file_extension}"}
                
        except Exception as e:
            return "", {"error": f"Extraction failed: {str(e)}", "file_path": str(file_path)}
    
    def _extract_pdf_text(self, file_path: Path, max_pages: int) -> Tuple[str, dict]:
        """Extract text from PDF file."""
        if not self.pdf_available:
            return "", {"error": "PyPDF2 not available for PDF extraction"}
        
        try:
            text_content = ""
            pages_processed = 0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                # Extract up to max_pages
                pages_to_extract = min(max_pages, total_pages)
                
                for page_num in range(pages_to_extract):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text_content += page_text + "\n\n"
                    pages_processed += 1
            
            metadata = {
                "file_type": "pdf",
                "total_pages": total_pages,
                "pages_processed": pages_processed,
                "pages_limit": max_pages,
                "text_length": len(text_content)
            }
            
            return text_content.strip(), metadata
            
        except Exception as e:
            return "", {"error": f"PDF extraction failed: {str(e)}"}
    
    def _extract_docx_text(self, file_path: Path, max_pages: int) -> Tuple[str, dict]:
        """Extract text from DOCX file."""
        if not self.docx_available:
            return "", {"error": "python-docx not available for DOCX extraction"}
        
        try:
            doc = docx.Document(file_path)
            text_content = ""
            paragraphs_processed = 0
            
            # DOCX doesn't have traditional pages, so we'll limit by paragraphs
            # Approximate: 50 paragraphs per "page" 
            max_paragraphs = max_pages * 50
            
            for paragraph in doc.paragraphs:
                if paragraphs_processed >= max_paragraphs:
                    break
                text_content += paragraph.text + "\n"
                paragraphs_processed += 1
            
            metadata = {
                "file_type": "docx",
                "total_paragraphs": len(doc.paragraphs),
                "paragraphs_processed": paragraphs_processed,
                "paragraphs_limit": max_paragraphs,
                "text_length": len(text_content)
            }
            
            return text_content.strip(), metadata
            
        except Exception as e:
            return "", {"error": f"DOCX extraction failed: {str(e)}"}
    
    def _extract_text_file(self, file_path: Path) -> Tuple[str, dict]:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    
                    metadata = {
                        "file_type": "text",
                        "encoding": encoding,
                        "text_length": len(text_content)
                    }
                    
                    return text_content, metadata
                    
                except UnicodeDecodeError:
                    continue
            
            return "", {"error": "Could not decode text file with any supported encoding"}
            
        except Exception as e:
            return "", {"error": f"Text file extraction failed: {str(e)}"}
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats."""
        formats = ['.txt', '.text']
        
        if self.pdf_available:
            formats.append('.pdf')
        if self.docx_available:
            formats.append('.docx')
            
        return formats
    
    def check_file_support(self, file_path: str) -> dict:
        """Check if file format is supported and available."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if not file_path.exists():
            return {"supported": False, "reason": "File not found"}
        
        if extension in ['.txt', '.text']:
            return {"supported": True, "library_required": None}
        elif extension == '.pdf':
            return {"supported": self.pdf_available, "library_required": "PyPDF2" if not self.pdf_available else None}
        elif extension == '.docx':
            return {"supported": self.docx_available, "library_required": "python-docx" if not self.docx_available else None}
        else:
            return {"supported": False, "reason": f"Unsupported format: {extension}"}