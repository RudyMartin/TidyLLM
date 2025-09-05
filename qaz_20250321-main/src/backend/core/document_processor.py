#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Processor Module

Handles real document processing including PDF text extraction,
document loading, and content analysis for QA criteria evaluation.
"""

import os
import sys
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
import re
from datetime import datetime

# Try to import PDF processing libraries
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logging.warning("pypdf not available - install with: pip install pypdf")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available - install with: pip install pdfplumber")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - install with: pip install python-docx")

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process documents for QA criteria analysis"""
    
    def __init__(self, input_dir: str = "input", output_dir: str = "output"):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
    def scan_input_directory(self) -> List[Path]:
        """Scan input directory for supported document types"""
        
        supported_extensions = ['.pdf', '.docx', '.txt', '.csv', '.xlsx']
        documents = []
        
        if not self.input_dir.exists():
            logger.warning(f"Input directory not found: {self.input_dir}")
            return documents
        
        for file_path in self.input_dir.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                documents.append(file_path)
                logger.info(f"Found document: {file_path.name}")
        
        logger.info(f"Found {len(documents)} documents in input directory")
        return documents
    
    def extract_text_from_pdf(self, pdf_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file using available libraries"""
        
        text_content = ""
        metadata = {
            'filename': pdf_path.name,
            'file_size': pdf_path.stat().st_size,
            'extraction_method': 'none',
            'pages': 0,
            'extraction_errors': []
        }
        
        try:
            # Try pdfplumber first (better text extraction)
            if PDFPLUMBER_AVAILABLE:
                logger.info(f"Extracting text from PDF using pdfplumber: {pdf_path.name}")
                text_content, metadata = self._extract_with_pdfplumber(pdf_path)
                metadata['extraction_method'] = 'pdfplumber'
                
            # Fallback to pypdf
            elif PYPDF_AVAILABLE:
                logger.info(f"Extracting text from PDF using pypdf: {pdf_path.name}")
                text_content, metadata = self._extract_with_pypdf(pdf_path)
                metadata['extraction_method'] = 'pypdf'
                
            else:
                logger.error("No PDF processing library available")
                metadata['extraction_errors'].append("No PDF library available")
                return text_content, metadata
                
        except Exception as e:
            logger.error(f"Error extracting text from PDF {pdf_path.name}: {e}")
            metadata['extraction_errors'].append(str(e))
        
        return text_content, metadata
    
    def _extract_with_pdfplumber(self, pdf_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text using pdfplumber"""
        
        text_content = ""
        metadata = {
            'filename': pdf_path.name,
            'file_size': pdf_path.stat().st_size,
            'pages': 0,
            'extraction_errors': []
        }
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num} ---\n"
                            text_content += page_text
                            text_content += "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {e}")
                        metadata['extraction_errors'].append(f"Page {page_num}: {str(e)}")
                
                # Extract metadata
                if pdf.metadata:
                    metadata.update({
                        'title': pdf.metadata.get('Title', ''),
                        'author': pdf.metadata.get('Author', ''),
                        'subject': pdf.metadata.get('Subject', ''),
                        'creator': pdf.metadata.get('Creator', ''),
                        'producer': pdf.metadata.get('Producer', ''),
                        'creation_date': pdf.metadata.get('CreationDate', ''),
                        'modification_date': pdf.metadata.get('ModDate', '')
                    })
        
        except Exception as e:
            logger.error(f"Error with pdfplumber extraction: {e}")
            metadata['extraction_errors'].append(str(e))
        
        return text_content, metadata
    
    def _extract_with_pypdf(self, pdf_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text using pypdf (modern replacement for PyPDF2)"""
        
        text_content = ""
        metadata = {
            'filename': pdf_path.name,
            'file_size': pdf_path.stat().st_size,
            'pages': 0,
            'extraction_errors': []
        }
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                metadata['pages'] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num} ---\n"
                            text_content += page_text
                            text_content += "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {e}")
                        metadata['extraction_errors'].append(f"Page {page_num}: {str(e)}")
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': pdf_reader.metadata.get('/CreationDate', ''),
                        'modification_date': pdf_reader.metadata.get('/ModDate', '')
                    })
        
        except Exception as e:
            logger.error(f"Error with pypdf extraction: {e}")
            metadata['extraction_errors'].append(str(e))
        
        return text_content, metadata
    
    def extract_text_from_docx(self, docx_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        
        text_content = ""
        metadata = {
            'filename': docx_path.name,
            'file_size': docx_path.stat().st_size,
            'extraction_method': 'docx',
            'extraction_errors': []
        }
        
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available for DOCX processing")
            metadata['extraction_errors'].append("python-docx library not available")
            return text_content, metadata
        
        try:
            logger.info(f"Extracting text from DOCX: {docx_path.name}")
            doc = docx.Document(docx_path)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract metadata
            metadata.update({
                'paragraphs': len(doc.paragraphs),
                'sections': len(doc.sections)
            })
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata.update({
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': props.created.isoformat() if props.created else '',
                    'modified': props.modified.isoformat() if props.modified else ''
                })
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path.name}: {e}")
            metadata['extraction_errors'].append(str(e))
        
        return text_content, metadata
    
    def extract_text_from_txt(self, txt_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from TXT file"""
        
        text_content = ""
        metadata = {
            'filename': txt_path.name,
            'file_size': txt_path.stat().st_size,
            'extraction_method': 'txt',
            'extraction_errors': []
        }
        
        try:
            logger.info(f"Extracting text from TXT: {txt_path.name}")
            
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                        metadata['encoding'] = encoding
                        break
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    logger.warning(f"Error reading with encoding {encoding}: {e}")
                    continue
            
            if not text_content:
                metadata['extraction_errors'].append("Could not read file with any encoding")
        
        except Exception as e:
            logger.error(f"Error extracting text from TXT {txt_path.name}: {e}")
            metadata['extraction_errors'].append(str(e))
        
        return text_content, metadata
    
    def process_document(self, file_path: Path) -> Dict[str, Any]:
        """Process a single document and extract content"""
        
        logger.info(f"Processing document: {file_path.name}")
        
        document_data = {
            'filename': file_path.name,
            'file_path': str(file_path),
            'file_type': file_path.suffix.lower(),
            'file_size': file_path.stat().st_size,
            'processed_at': datetime.now().isoformat(),
            'content': '',
            'metadata': {},
            'extraction_errors': []
        }
        
        try:
            # Extract text based on file type
            if file_path.suffix.lower() == '.pdf':
                content, metadata = self.extract_text_from_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                content, metadata = self.extract_text_from_docx(file_path)
            elif file_path.suffix.lower() == '.txt':
                content, metadata = self.extract_text_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_path.suffix}")
                document_data['extraction_errors'].append(f"Unsupported file type: {file_path.suffix}")
                return document_data
            
            document_data['content'] = content
            document_data['metadata'] = metadata
            document_data['extraction_errors'] = metadata.get('extraction_errors', [])
            
            logger.info(f"Successfully processed {file_path.name}: {len(content)} characters")
        
        except Exception as e:
            logger.error(f"Error processing document {file_path.name}: {e}")
            document_data['extraction_errors'].append(str(e))
        
        return document_data
    
    def process_all_documents(self) -> List[Dict[str, Any]]:
        """Process all documents in the input directory"""
        
        logger.info("Starting document processing...")
        
        # Scan for documents
        document_files = self.scan_input_directory()
        
        if not document_files:
            logger.warning("No documents found in input directory")
            return []
        
        # Process each document
        processed_documents = []
        
        for file_path in document_files:
            document_data = self.process_document(file_path)
            processed_documents.append(document_data)
        
        # Log summary
        successful = sum(1 for doc in processed_documents if not doc['extraction_errors'])
        failed = len(processed_documents) - successful
        
        logger.info(f"Document processing complete: {successful} successful, {failed} failed")
        
        return processed_documents
    
    def extract_metadata_fields(self, documents: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Extract metadata fields from processed documents"""
        
        logger.info("Extracting metadata fields from documents...")
        
        extracted_fields = {
            'review_id': 'REV00000',  # Default
            'model_type': 'Unknown',
            'risk_tier': 'Medium',
            'model_id': 'Unknown',
            'model_name': 'Unknown Model',
            'version': '1.0.0',
            'authors': ['Unknown'],
            'date': datetime.now().strftime("%m-%d-%Y"),
            'validation_type': 'Standard Review',
            'reviewer_name': 'Unknown',
            'team_num': 'Unknown',
            'process_name': 'QA Validation Review'
        }
        
        # Search patterns for different fields
        patterns = {
            'review_id': [
                r'Review ID[:\s]*([A-Z]{3}\d{5})',
                r'Review[:\s]*([A-Z]{3}\d{5})',
                r'([A-Z]{3}\d{5})',
                r'Review ID[:\s]*(\d{5})',
            ],
            'model_type': [
                r'Model Type[:\s]*([^\n\r]+)',
                r'Type[:\s]*([^\n\r]+)',
            ],
            'risk_tier': [
                r'Risk Tier[:\s]*([^\n\r]+)',
                r'Risk Level[:\s]*([^\n\r]+)',
            ],
            'model_id': [
                r'Model ID[:\s]*([^\n\r]+)',
                r'ID[:\s]*([^\n\r]+)',
            ],
            'model_name': [
                r'Model Name[:\s]*([^\n\r]+)',
                r'Model[:\s]*([^\n\r]+)',
            ],
            'version': [
                r'Version[:\s]*([^\n\r]+)',
                r'v\.([^\n\r]+)',
            ],
            'authors': [
                r'Authors?[:\s]*([^\n\r]+)',
                r'Author[:\s]*([^\n\r]+)',
            ],
            'date': [
                r'Date[:\s]*([^\n\r]+)',
                r'(\d{2}-\d{2}-\d{4})',
            ],
            'validation_type': [
                r'Validation Type[:\s]*([^\n\r]+)',
                r'Type[:\s]*([^\n\r]+)',
            ]
        }
        
        # Search through all documents
        for document in documents:
            content = document.get('content', '')
            
            for field, field_patterns in patterns.items():
                for pattern in field_patterns:
                    match = re.search(pattern, content, re.IGNORECASE)
                    if match:
                        value = match.group(1).strip()
                        
                        # Special handling for different fields
                        if field == 'review_id':
                            if not value.startswith('REV'):
                                value = f"REV{value.zfill(5)}"
                        elif field == 'authors':
                            # Split authors by comma or semicolon
                            authors = re.split(r'[,;]', value)
                            authors = [author.strip() for author in authors if author.strip()]
                            if authors:
                                extracted_fields[field] = authors
                        elif field == 'date':
                            # Try to parse and format date
                            try:
                                # Add date parsing logic here if needed
                                pass
                            except:
                                pass
                        
                        if field != 'authors':  # Already handled above
                            extracted_fields[field] = value
                        
                        logger.info(f"Found {field}: {value}")
                        break  # Use first match for each field
        
        logger.info(f"Extracted {len(extracted_fields)} metadata fields")
        return extracted_fields


def main():
    """Test the document processor"""
    
    print("🚀 Testing Document Processor")
    print("=" * 40)
    
    processor = DocumentProcessor()
    
    # Process all documents
    documents = processor.process_all_documents()
    
    if documents:
        print(f"\n📋 Processed {len(documents)} documents:")
        for doc in documents:
            print(f"  • {doc['filename']}: {len(doc['content'])} characters")
            if doc['extraction_errors']:
                print(f"    ⚠️ Errors: {doc['extraction_errors']}")
        
        # Extract metadata
        metadata = processor.extract_metadata_fields(documents)
        print(f"\n📊 Extracted Metadata:")
        for key, value in metadata.items():
            print(f"  • {key}: {value}")
    else:
        print("❌ No documents processed")


if __name__ == "__main__":
    main()
