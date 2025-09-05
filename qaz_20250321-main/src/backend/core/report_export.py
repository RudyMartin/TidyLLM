
import streamlit as st
import os
import tempfile
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from datetime import datetime
try:
    from core.config import CONFIG
except ImportError:
    # Fallback for deployment environment
    from backend.core.config import CONFIG
import boto3

def export_dashboard_to_pdf(topic_fig, cm_fig, topic_scores_df, winners_df, title="VectorQA Evaluation Report"):
    with tempfile.TemporaryDirectory() as tmpdir:
        topic_img = os.path.join(tmpdir, "topic_accuracy.png")
        cm_img = os.path.join(tmpdir, "confusion_matrix.png")
        topic_fig.savefig(topic_img)
        cm_fig.savefig(cm_img)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=14)
        pdf.cell(200, 10, title, ln=True, align='C')
        pdf.set_font("Arial", size=10)
        pdf.cell(200, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align='C')
        pdf.ln(10)

        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, "Accuracy by Topic", ln=True)
        pdf.image(topic_img, x=10, y=40, w=180)
        pdf.ln(90)

        pdf.cell(200, 10, "Confusion Matrix", ln=True)
        pdf.image(cm_img, x=10, y=130, w=180)
        pdf.add_page()

        pdf.set_font("Arial", size=11)
        pdf.cell(200, 10, "Topic Accuracy Table", ln=True)
        for i, row in topic_scores_df.iterrows():
            pdf.cell(0, 10, f"{row['topic']}: {row['accuracy']:.2f}", ln=True)

        pdf.ln(10)
        pdf.cell(200, 10, "Winner by Topic", ln=True)
        for i, row in winners_df.iterrows():
            pdf.cell(0, 10, f"{row['topic']} - {row['model']}: {row['accuracy']:.2f}", ln=True)

        pdf_output = os.path.join(tmpdir, "evaluation_report.pdf")
        pdf.output(pdf_output)

        return open(pdf_output, "rb").read()

def export_dashboard_to_docx(topic_scores_df, winners_df, title="VectorQA Evaluation Report"):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("Topic Accuracy Table", level=2)
    for i, row in topic_scores_df.iterrows():
        doc.add_paragraph(f"{row['topic']}: {row['accuracy']:.2f}")

    doc.add_heading("Winner by Topic", level=2)
    for i, row in winners_df.iterrows():
        doc.add_paragraph(f"{row['topic']} - {row['model']}: {row['accuracy']:.2f}")

    tmpfile = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmpfile.name)
    return open(tmpfile.name, "rb").read()

def export_dashboard_to_html(topic_scores_df, winners_df, title="VectorQA Evaluation Report"):
    html = f"<h1>{title}</h1><p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>"
    html += "<h2>Topic Accuracy Table</h2><ul>"
    for i, row in topic_scores_df.iterrows():
        html += f"<li>{row['topic']}: {row['accuracy']:.2f}</li>"
    html += "</ul><h2>Winner by Topic</h2><ul>"
    for i, row in winners_df.iterrows():
        html += f"<li>{row['topic']} - {row['model']}: {row['accuracy']:.2f}</li>"
    html += "</ul>"

    tmpfile = tempfile.NamedTemporaryFile(delete=False, suffix=".html")
    with open(tmpfile.name, "w") as f:
        f.write(html)
    return open(tmpfile.name, "rb").read()

def optionally_upload_to_s3(content_bytes, filename, content_type="application/pdf"):
    if "bucket_name" not in CONFIG:
        return None
    bucket = CONFIG["bucket_name"]
    s3_path = f"reports/{filename}"
    s3 = boto3.client("s3")
    s3.put_object(Bucket=bucket, Key=s3_path, Body=content_bytes, ContentType=content_type)
    return f"s3://{bucket}/{s3_path}"
