"""
MVR Document Processing Service - Model Validation Report Analysis
=================================================================
🔧 UTILITY SERVICE - NOT A CORE GATEWAY
This is a specialized document processing service, not part of the main gateway workflow.

Purpose: Process Model Validation Reports (regulatory compliance documents) for:
- Banking/Finance: Basel III compliance validation
- Healthcare: FDA model validation requirements  
- Insurance: Regulatory model validation standards
"""

from .base_gateway import BaseGateway
from typing import Dict, Any, List, Optional
import json
from pathlib import Path
import PyPDF2
import docx
from datetime import datetime


class MVRAnalysisGateway(BaseGateway):
    """
    Specialized gateway for MVR document analysis and report generation.
    
    This gateway coordinates between:
    - Document extraction (PDF/DOCX)
    - Prompt template loading
    - LLM processing
    - Report generation
    """
    
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.prompt_dir = Path("qaz_20250321-main/src/assets/prompts/favorites")
        self.report_cache = {}
        
    def process(self, document_path: str, report_type: str, **kwargs) -> Dict[str, Any]:
        """
        Process an MVR document for a specific report type.
        
        Args:
            document_path: Path to the MVR document
            report_type: Type of report (compliance, intelligence, knowledge)
            **kwargs: Additional processing parameters
            
        Returns:
            Processed analysis results
        """
        # Extract document content
        content = self._extract_document_content(document_path)
        
        # Load appropriate prompt template
        prompt = self._load_prompt_template(report_type)
        
        # Process through appropriate analysis pipeline
        if report_type == "compliance":
            result = self._process_compliance_analysis(content, prompt, **kwargs)
        elif report_type == "intelligence":
            result = self._process_intelligence_analysis(content, prompt, **kwargs)
        elif report_type == "knowledge":
            result = self._process_knowledge_extraction(content, prompt, **kwargs)
        else:
            raise ValueError(f"Unknown report type: {report_type}")
        
        return result
    
    def _extract_document_content(self, document_path: str) -> Dict[str, Any]:
        """Extract content from PDF or DOCX document."""
        path = Path(document_path)
        content = {
            "filename": path.name,
            "format": path.suffix.lower(),
            "text": "",
            "toc": [],
            "figures": [],
            "tables": [],
            "references": []
        }
        
        if path.suffix.lower() == ".pdf":
            content.update(self._extract_pdf_content(path))
        elif path.suffix.lower() in [".docx", ".doc"]:
            content.update(self._extract_docx_content(path))
        else:
            raise ValueError(f"Unsupported document format: {path.suffix}")
        
        return content
    
    def _extract_pdf_content(self, path: Path) -> Dict[str, Any]:
        """Extract content from PDF file."""
        content = {"text": "", "pages": 0}
        
        try:
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content["pages"] = len(pdf_reader.pages)
                
                text_parts = []
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                
                content["text"] = "\n\n".join(text_parts)
                
                # Extract TOC if available
                if pdf_reader.outline:
                    content["toc"] = self._extract_pdf_toc(pdf_reader.outline)
                    
        except Exception as e:
            print(f"Error extracting PDF content: {e}")
        
        return content
    
    def _extract_pdf_toc(self, outline, level=0) -> List[Dict]:
        """Recursively extract PDF table of contents."""
        toc = []
        for item in outline:
            if isinstance(item, list):
                # Nested items
                toc.extend(self._extract_pdf_toc(item, level + 1))
            else:
                toc.append({
                    "title": item.title if hasattr(item, 'title') else str(item),
                    "level": level,
                    "page": item.page.idnum if hasattr(item, 'page') else None
                })
        return toc
    
    def _extract_docx_content(self, path: Path) -> Dict[str, Any]:
        """Extract content from DOCX file."""
        content = {"text": "", "paragraphs": 0}
        
        try:
            doc = docx.Document(path)
            
            # Extract text
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            content["text"] = "\n\n".join(text_parts)
            content["paragraphs"] = len(text_parts)
            
            # Extract tables
            content["tables"] = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content["tables"].append(table_data)
                
        except Exception as e:
            print(f"Error extracting DOCX content: {e}")
        
        return content
    
    def _load_prompt_template(self, report_type: str) -> str:
        """Load the appropriate prompt template."""
        template_map = {
            "compliance": "JB_Overview_Prompt.md",
            "intelligence": "comprehensive_whitepaper_analysis.md",
            "knowledge": "toc_extraction_prompt.md"
        }
        
        template_file = self.prompt_dir / template_map.get(report_type, "")
        
        if template_file.exists():
            with open(template_file, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            print(f"Warning: Template not found at {template_file}")
            return self._get_default_prompt(report_type)
    
    def _get_default_prompt(self, report_type: str) -> str:
        """Get default prompt if template is not found."""
        defaults = {
            "compliance": "Analyze this Model Validation Report for compliance with regulatory standards.",
            "intelligence": "Extract key insights and intelligence from this document.",
            "knowledge": "Extract the table of contents and identify references for knowledge expansion."
        }
        return defaults.get(report_type, "Analyze this document.")
    
    def _process_compliance_analysis(self, content: Dict, prompt: str, **kwargs) -> Dict[str, Any]:
        """Process compliance analysis using the MVR peer review prompt."""
        result = {
            "type": "compliance",
            "timestamp": datetime.now().isoformat(),
            "sections": []
        }
        
        # Parse document into sections
        sections = self._parse_document_sections(content["text"])
        
        # Analyze each section for compliance
        for section in sections:
            analysis = {
                "section_id": section["id"],
                "section_title": section["title"],
                "mvs_requirements": [],
                "review_narrative": "",
                "contradiction_summary": "None",
                "peer_review_challenge": "",
                "conclusion": "✅",
                "confidence_score": "Highly Confident",
                "defect_type": "N/A"
            }
            
            # This would normally call the LLM with the prompt
            # For now, we'll add placeholder logic
            result["sections"].append(analysis)
        
        return result
    
    def _process_intelligence_analysis(self, content: Dict, prompt: str, **kwargs) -> Dict[str, Any]:
        """Process document intelligence analysis."""
        result = {
            "type": "intelligence",
            "timestamp": datetime.now().isoformat(),
            "document_stats": {
                "pages": content.get("pages", 0),
                "paragraphs": content.get("paragraphs", 0),
                "tables": len(content.get("tables", [])),
                "toc_sections": len(content.get("toc", []))
            },
            "key_findings": [],
            "section_summaries": [],
            "visual_elements": []
        }
        
        # Extract key sections for summarization
        sections = self._parse_document_sections(content["text"])
        
        for section in sections[:10]:  # Limit to first 10 sections
            summary = {
                "section": section["title"],
                "summary": f"Summary of {section['title']}",
                "word_count": len(section.get("text", "").split())
            }
            result["section_summaries"].append(summary)
        
        return result
    
    def _process_knowledge_extraction(self, content: Dict, prompt: str, **kwargs) -> Dict[str, Any]:
        """Process knowledge extraction and reference discovery."""
        result = {
            "type": "knowledge",
            "timestamp": datetime.now().isoformat(),
            "toc": content.get("toc", []),
            "discovered_papers": [],
            "references": []
        }
        
        # Extract references from text
        references = self._extract_references(content["text"])
        
        for ref in references:
            paper = {
                "title": ref.get("title", ""),
                "authors": ref.get("authors", ""),
                "year": ref.get("year", ""),
                "source": "unknown",
                "url": "",
                "relevance_score": 0.0,
                "availability": "unknown",
                "domain": "unknown",
                "validation_status": "pending"
            }
            result["discovered_papers"].append(paper)
        
        return result
    
    def _parse_document_sections(self, text: str) -> List[Dict]:
        """Parse document text into sections."""
        sections = []
        lines = text.split('\n')
        
        current_section = None
        section_id = 0
        
        for line in lines:
            # Simple heuristic: lines that look like headers
            if line.strip() and (
                line.isupper() or 
                line.startswith(('1.', '2.', '3.', '4.', '5.')) or
                any(line.startswith(f"{i}.") for i in range(1, 10))
            ):
                if current_section:
                    sections.append(current_section)
                
                section_id += 1
                current_section = {
                    "id": f"{section_id}",
                    "title": line.strip(),
                    "text": ""
                }
            elif current_section:
                current_section["text"] += line + "\n"
        
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _extract_references(self, text: str) -> List[Dict]:
        """Extract references from document text."""
        references = []
        
        # Simple pattern matching for references
        # This would be more sophisticated in production
        lines = text.split('\n')
        
        for line in lines:
            if 'et al.' in line or '(20' in line or '(19' in line:
                references.append({
                    "raw_text": line.strip(),
                    "title": "",
                    "authors": "",
                    "year": ""
                })
        
        return references[:20]  # Limit to first 20 references